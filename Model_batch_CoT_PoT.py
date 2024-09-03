from docx import Document
from prompts.CoT_good_prompts import CoT_en_prompts
import json
from datetime import datetime
import torch
import os
from utils import chat_utils, dolphin_chat_utils, neural_chat_utils, openorca_chat_utils, prompts

from utils.chat_utils import get_gpu_memory
from utils.docx_utils import add_para, add_title, add_dictionary, add_para_highlight
from utils.intercept_error import convert_full_text, check_indent_and_run
from utils.logger_config import LoggerConfig

os.environ["TOKENIZERS_PARALLELISM"] = "false"
torch.cuda.empty_cache()
print(get_gpu_memory())
date = datetime.now().date()
now_time = datetime.now().time()

model_dir = "/home/ligy/zmb_workspace/model/"
auto_dir = "prompts/template/Auto/"
save_json_dir = "outputs/auto/json/"

log_dir = 'outputs/auto_cot3_pot/log/{0}'.format(date)
os.makedirs(log_dir, exist_ok=True)
logger = LoggerConfig(log_file='{0}/{1}.log'.format(log_dir, now_time)).logger

model_device1 = 'auto'
model_tokenizer_device = 'cuda'
model1_name = "dolphin-2.2.1-mistral-7b"
model1_path = model_dir + model1_name
model1, tokenizer1 = dolphin_chat_utils.load_model(model1_path, model_device1)
logger.info(get_gpu_memory())

doc = Document()
doc.add_heading('已采纳题目结果', 0)
dataset_dir = "dataset/benchmark/"
adopt_json_file = dataset_dir + "adopt.json"
save_docx_dir = "outputs/auto_cot3_pot/docx/"
save_docx_file = save_docx_dir + "results.docx"
os.makedirs(save_docx_dir, exist_ok=True)
doc.save(save_docx_file)

role_path = "prompts/model_role/financial_analyst_assistant_sample"
system_m = open(role_path, 'r', encoding='utf-8').read().replace("\n", " ")


def main():
    q_sum = 0
    q_correct = 0
    adopt_json = json.load(open(adopt_json_file, 'r'))
    for dataset_name, adopt_q_list in adopt_json.items():
        q_num = 0
        file_in = open(dataset_dir + dataset_name + ".txt", 'r', encoding='utf-8')
        for query in file_in:
            q_num += 1
            if q_num in adopt_q_list:
                try:
                    add_doc = Document(save_docx_file)
                    q_sum += 1
                    history = []
                    query_list = query.strip().split("###")
                    question = query_list.pop(-1)
                    question = convert_full_text(question).strip()
                    logger.info("{0}: 题目{1}: {2}".format(dataset_name, q_num, question))
                    answer_list = query_list

                    # CoT
                    CoT1_simple = ("Below is a question about financial calculation.\nQuestion: {question}\n"
                                   "Let's solve the question step by step.")
                    CoT2_heuristic = ("Please first list all the numerical information in the question. "
                                      "Focus exclusively on the numbers given in the question, assigning each a "
                                      "unique variable."
                                      " Then list all calculation formulas.\nQuestion: {question}\nAnswer:")
                    CoT3 = ("Below is a question about financial calculation.\nQuestion: {question}\n"
                            "Please list all variables and required calculation formulas in the question."
                            "Then let's solve the question step by step.")
                    CoT_input = CoT3.format(question=question)
                    history, _, _ = dolphin_chat_utils.chat(
                        history, system_m, model1, tokenizer1, model1_name, CoT_input, logger, model_tokenizer_device,
                        False)

                    # PoT
                    PoT_input = "Let's write a Python program."
                    history, PoT_output, input_all = dolphin_chat_utils.chat(
                        history, system_m, model1, tokenizer1, model1_name, PoT_input, logger, model_tokenizer_device,
                        False)
                    output_code = prompts.extract_code(PoT_output)
                    PoT_answer = check_indent_and_run(output_code)

                    logger.info("answer_list: {0}\npot_result: {1}\n\n".format(answer_list, PoT_answer))
                    logger.info("### PoT_input\n{0}\n\n### PoT_output\n{1}\n\n".format(input_all, PoT_output))
                    if all(elem.strip() in PoT_answer.split(",") for elem in answer_list):
                        add_title(add_doc, "\n\n{0}: 题目{1}, 答案正确".format(dataset_name, q_num))
                        add_para(add_doc, question)
                        q_correct += 1
                        add_para(add_doc, "### PoT_input\n{0}\n\n### PoT_output\n{1}\n\n".format(input_all, PoT_output))
                    else:
                        add_title(add_doc, "\n\n{0}: 题目{1}, 答案错误".format(dataset_name, q_num))
                        add_para(add_doc, question)
                    add_para_highlight(add_doc, "answer_list: {0}\npot_result: {1}\n\n".format(answer_list, PoT_answer))
                    add_para(add_doc, "### PoT_input\n{0}\n\n### PoT_output\n{1}\n\n".format(input_all, PoT_output))

                    add_doc.save(save_docx_file)
                    logger.info("### q_sum: {0}, q_correct: {1}, accuracy：{2}% ###".format(
                        q_sum, q_correct, (q_correct / q_sum) * 100))
                except Exception as e:
                    logger.info("{0}, q_{1}, New Error: {2}".format(dataset_name, q_num, str(e)))
                    continue


if __name__ == '__main__':
    main()
