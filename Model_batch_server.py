import queue
import requests
import threading

from docx import Document

from Create_json import create_json
from utils.cross_models_utils import *
from Create_tree_paths import *
import json
from datetime import datetime
import torch
import os
from utils import chat_utils, dolphin_chat_utils, neural_chat_utils, openorca_chat_utils

from utils.chat_utils import get_gpu_memory
from utils.docx_utils import add_para, add_title, add_dictionary, add_para_highlight
from utils.intercept_error import convert_full_text
from utils.logger_config import LoggerConfig
from utils.sympy_for_paths import paths_to_sympy, sympy_check, list_to_sympy

torch.cuda.empty_cache()
print(get_gpu_memory())
date = datetime.now().date()
now_time = datetime.now().time()

model_dir = "/mnt/zmb/zmb_workspace/model/"
auto_dir = "prompts/template/Auto/"

model1_name = "dolphin-2.2.1-mistral-7b"
model2_name = "dolphin-2.6-mistral-7b-dpo"
dolphin_2_6_server = "http://192.168.0.105:8080/"

make_dir = 'outputs/auto_server/dolphin-2.2.1+dolphin-2.6'
log_dir = make_dir + '/log/{0}'.format(date)
os.makedirs(log_dir, exist_ok=True)
logger = LoggerConfig(log_file='{0}/{1}.log'.format(log_dir, now_time)).logger

result = requests.get(dolphin_2_6_server).json()
logger.info("Get from {0}({1}): {2}".format(model2_name, dolphin_2_6_server, result))
model_device1 = 'auto'
model_tokenizer_device = 'cuda'
model1_path = model_dir + model1_name
model1, tokenizer1 = dolphin_chat_utils.load_model(model1_path, model_device1)
logger.info(get_gpu_memory())

doc = Document()
doc.add_heading('已采纳题目结果', 0)
dataset_dir = 'dataset/benchmark/'
adopt_json_file = dataset_dir + "adopt.json"
save_docx_dir = make_dir + '/docx/'
save_docx_file = save_docx_dir + "results.docx"
save_docx_log_dir = make_dir + '/docx/log/'
save_json_dir = make_dir + '/json/'
os.makedirs(save_docx_dir, exist_ok=True)
doc.save(save_docx_file)


def main():
    q_sum = 0
    q_correct = 0
    adopt_json = json.load(open(adopt_json_file, 'r'))
    for dataset_name, adopt_q_list in adopt_json.items():
        q_num = 0
        file_in = open(dataset_dir + dataset_name + ".txt", 'r', encoding='utf-8')
        for query in file_in:
            q_num += 1
            if q_num in adopt_q_list:
                try:
                    add_doc = Document(save_docx_file)
                    log_doc = Document()
                    log_doc.add_heading("{0}: 题目{1}, 多模型多角色路径筛选".format(dataset_name, q_num), 0)
                    q_sum += 1
                    query_list = query.strip().split("###")
                    question = query_list.pop(-1)
                    question = convert_full_text(question)
                    logger.info("{0}: 题目{1}: {2}".format(dataset_name, q_num, question))
                    save_json_file = save_json_dir + "{0}/q_{1}.json".format(dataset_name, q_num)
                    if os.path.exists(save_json_file):
                        q_json = json.load(open(save_json_file, 'r', encoding='utf-8'))
                    else:
                        json_log_doc = Document()
                        json_log_doc.add_heading("{0}: 题目{1}, 单模型抽取JSON日志".format(dataset_name, q_num), 0)
                        q_json = create_json(model1_name, question, model1, tokenizer1, model_tokenizer_device, logger,
                                             save_json_file, json_log_doc)
                        os.makedirs(save_docx_log_dir + dataset_name, exist_ok=True)
                        json_log_doc.save(save_docx_log_dir + "{0}/{0}_q_{1}_json.docx".format(dataset_name, q_num))
                    answer_list = query_list

                    if not check_json_completeness(q_json):
                        logger.info("Json is incomplete !!!")
                        add_title(add_doc, "\n\n{0}: 题目{1}, 答案错误".format(dataset_name, q_num))
                        add_para(add_doc, question)
                        add_para_highlight(add_doc, "Json is incomplete !!!\nJSON:")
                        add_dictionary(add_doc, q_json, 1)
                        add_doc.save(save_docx_file)
                        continue
                    q_json_v = {"Target Variable(s)": q_json["Target Variable(s)"],
                                "Input Variable(s)": q_json["Input Variable(s)"]}
                    q_json_f = q_json["List all formula(s)/equation(s)"]

                    sympy_result_dict = {}
                    origin_formula_list = []
                    for key, value in q_json_f.items():
                        formula = check_and_correct_parentheses(value)
                        formula = replace_multiple_spaces_with_underscores(formula)
                        origin_formula_list.append(formula)
                    origin_variables_set = set()
                    for formula in origin_formula_list:
                        origin_variables_set.update(extract_variables(formula))

                    formula_list = sympy_check(origin_formula_list, origin_variables_set, logger)

                    variables_set = set()
                    for formula in formula_list:
                        variables_set.update(extract_variables(formula))
                    input_values_dict = {}
                    for variable in q_json["Input Variable(s)"]:
                        if str_is_number(variable["Numerical Value"]):
                            input_values_dict[variable["Variable Name"].replace("'", "")] = variable["Numerical Value"]
                    target_builder = create_tree_and_paths(logger, q_json, formula_list)
                    print_docx = ""
                    for tree in target_builder:
                        if tree.key_index > 1:
                            logger.info("Formula_list: \n{0}".format(formula_list))
                            logger.info("Variables_set: \n{0}".format(variables_set))
                            logger.info("Input_values_dict: \n{0}".format(input_values_dict))

                            target_v = tree.root.value
                            all_formula_paths = extract_formula_paths(tree.root)
                            logger.info("Sympy Paths Count：{0}".format(len(all_formula_paths)))
                            if len(all_formula_paths) > 200:
                                error_log = "Target Variable: {0}. Sympy Paths Count：{1} > 200, too many !!!\n".format(
                                    target_v, len(all_formula_paths))
                                logger.info(error_log)
                                print_docx += error_log
                                continue

                            result_list = []
                            for formula_path in all_formula_paths:
                                logger.info("===== Path =====")
                                result = paths_to_sympy(formula_path, variables_set, input_values_dict, target_v, logger)
                                logger.info("Result: {0}".format(result))
                                result_list.append(result)

                            solution_paths = extract_paths(tree.root, input_values_dict)
                            logger.info("Solution_Paths: \n{0}".format(print_paths(solution_paths)))
                            logger.info("Result List: \n{0}".format(result_list))

                            paths_count = len(solution_paths)
                            remove_sympy_not_num(result_list, solution_paths, all_formula_paths)
                            # remove_sympy_error_none(result_list, solution_paths, all_formula_paths)

                            if 0 in result_list and check_0_unreasonable(target_v, q_json_v, log_doc):
                                remove_sympy_0(result_list, solution_paths, all_formula_paths)

                            if contains_positive_and_negative(result_list) and check_negative_unreasonable(
                                    target_v, q_json_v, log_doc):
                                remove_sympy_negative(result_list, solution_paths, all_formula_paths)

                            sympy_num_paths_count = len(solution_paths)

                            solution_paths_str = print_paths(solution_paths)
                            logger.info("Solution_Paths after remove: \n{0}".format(solution_paths_str))
                            log_str = ("Target Variable: {0}, Paths Count：{1}, Sympy Num Paths Count：{2}\nResult "
                                       "List: {3}"
                                       .format(target_v, paths_count, sympy_num_paths_count, result_list))
                            logger.info(log_str)
                            print_docx += log_str
                            if len(solution_paths) > 26:
                                print_str = ("Target Variable: {0}. Sympy num Paths Count：{1}, too many !!!".
                                             format(target_v, sympy_num_paths_count))
                                logger.info(print_str)
                                print_docx += "\n{0}\n".format(print_str)
                            elif len(solution_paths) == 0:
                                print_str = "Target Variable: {0}. Have no Sympy Number Result !!!".format(target_v)
                                logger.info(print_str)
                                print_docx += "\n{0}\n".format(print_str)
                            else:
                                letter_paths_dict = letter_paths(solution_paths)
                                sympy_result, best_choice = model_cross(letter_paths_dict, question, q_json_v,
                                                                        all_formula_paths, variables_set,
                                                                        input_values_dict, target_v, log_doc)
                                input_values_dict[tree.root.value] = sympy_result
                                sympy_result_dict[target_v] = sympy_result
                                print_docx += ("\n\nTarget Variable: {0}\n\nBest_path:\n{1}\nSympy_result: {2}\n".
                                               format(target_v, best_choice, sympy_result))
                                logger.info("Target Variable: {0}, Final Sympy Result: {1}".format(target_v, sympy_result))
                                add_para_highlight(log_doc, "\n\nBest_path:\n{0}\nSympy_result: {1}\n".format(best_choice, sympy_result))
                        else:
                            logger.info("Can not create a tree. Use list_to_sympy():")
                            print_docx += "\n\nTarget Variable: {0}, Can not create a tree !!!\n".format(tree.root.value)
                            logger.info("Formula_list: {0}".format(formula_list))
                            # all_solutions = list_to_sympy(formula_list, variables_set, input_values_dict, logger)
                            # logger.info("All_solutions: {0}".format(all_solutions))

                    logger.info("sympy_result: {0}, answer_list: {1}".format(list(sympy_result_dict.values()), answer_list))
                    if all(elem in list(sympy_result_dict.values()) for elem in answer_list):
                        add_title(add_doc, "\n\n{0}: 题目{1}, 答案正确".format(dataset_name, q_num))
                        add_para(add_doc, question)
                        q_correct += 1
                    else:
                        add_title(add_doc, "\n\n{0}: 题目{1}, 答案错误".format(dataset_name, q_num))
                        add_para(add_doc, question)
                        add_para(add_doc, "answer_list: {0}\nsympy_result_dict: {1}\n\nJSON:".format(
                            answer_list, sympy_result_dict))
                        add_dictionary(add_doc, q_json, 1)
                    add_para_highlight(add_doc, print_docx)
                    add_doc.save(save_docx_file)
                    log_doc.save(save_docx_log_dir + "{0}/{0}_q_{1}_choose.docx".format(dataset_name, q_num))
                    logger.info("### q_sum: {0}, q_correct: {1}, accuracy：{2}% ###".format(
                        q_sum, q_correct, (q_correct / q_sum) * 100))
                except Exception as e:
                    logger.info("{0}, q_{1}, New Error: {2}".format(dataset_name, q_num, str(e)))
                    continue


def model_cross(letter_paths_dict, question, q_json_v, all_formula_paths, variables_set, input_values_dict, target_v, log_doc):
    logger.info('run model_cross(): {0}, {1}'.format(model1_name, model2_name))
    add_title(log_doc, "\n##### 选择最优路径 #####:\n")

    if len(letter_paths_dict) > 1:
        pairs, best_choice_index, best_choice = binary_choice_optimization(letter_paths_dict, question, q_json_v, log_doc)
        formula_path = all_formula_paths[best_choice_index]
    else:
        formula_path = all_formula_paths[0]
        best_choice = list(letter_paths_dict.items())[0]
    return path_step_check(formula_path, variables_set, input_values_dict, target_v), best_choice[1]


def path_step_check(formula_path, variables_set, input_values_dict, target_v):
    sympy_result = paths_to_sympy(formula_path, variables_set, input_values_dict, target_v, logger)
    logger.info("### Optimal Path, sympy_result: {0}, formula_path: {1}".format(sympy_result, formula_path))
    return sympy_result


def binary_choice_optimization(letter_paths_dict, question, q_json_v, log_doc):
    original_pairs = list(letter_paths_dict.items())
    pairs = original_pairs.copy()
    while len(pairs) > 1:
        for i in range(0, len(pairs) - 1, 2):
            chosen_pair = choose_one_path(pairs[i], pairs[i + 1], question, q_json_v, log_doc)
            logger.info(f"Chosen_pair：{chosen_pair}")
            pairs.remove(pairs[i] if chosen_pair == pairs[i + 1] else pairs[i + 1])
            break
    best_choice = pairs[0]
    best_choice_index = original_pairs.index(best_choice)
    return pairs, best_choice_index, best_choice


def choose_one_path(path1, path2, question, q_json_v, log_doc):
    history_model1 = []
    history_model2 = []
    role_path = "prompts/model_role/debater"
    system_m = open(role_path, 'r', encoding='utf-8').read().replace("\n", " ")

    solution_2_diff = open(auto_dir + "cross/solution_2_diff.txt", 'r').read().format(
        Question=question, variable='{0}'.format(q_json_v), solution1=path1[1],
        solution2=path2[1], solution1_n=path1[0], solution2_n=path2[0])
    solution_2_choose = open(auto_dir + "cross/solution_2_choose_3.txt", 'r').read()
    solution_2_json = open(auto_dir + "cross/solution_2_json.txt", 'r').read()

    # model2
    def model2_func(history_model2, system_m, solution_2_diff, solution_2_choose, solution_2_json, result_queue):
        history_model2, _, _ = post_server(history_model2, system_m, solution_2_diff, False)
        history_model2, _, _ = post_server(history_model2, system_m, solution_2_choose, False)
        history_model2, model2_output_json, input_all = post_server(history_model2, system_m, solution_2_json, True)
        result_queue.put(history_model2)
        result_queue.put(model2_output_json)
        result_queue.put(input_all)

    result_queue = queue.Queue()
    model2_thread = threading.Thread(target=model2_func, args=(history_model2, system_m, solution_2_diff,
                                                               solution_2_choose, solution_2_json, result_queue))
    model2_thread.start()

    # model1
    model_name = model1_name
    history_model1, _, _ = dolphin_chat_utils.chat(history_model1, system_m, model1, tokenizer1, model_name,
                                                   solution_2_diff, logger, model_tokenizer_device, False)
    history_model1, _, _ = dolphin_chat_utils.chat(history_model1, system_m, model1, tokenizer1, model_name,
                                                   solution_2_choose, logger, model_tokenizer_device, False)
    history_model1, model1_output_json, input_all = dolphin_chat_utils.chat(history_model1, system_m, model1,
                                                                            tokenizer1,
                                                                            model_name, solution_2_json, logger,
                                                                            model_tokenizer_device, True)
    logger.info("\n##### {0} Discuss Input_all #####\n{1}".format(model1_name, input_all))
    logger.info('\n##### {0} Discuss Answer #####\n{1}\n\n'.format(model1_name, model1_output_json))
    add_para(log_doc, "\n##### {0} Discuss Input_all #####\n{1}".format(model1_name, input_all))
    add_para_highlight(log_doc, '\n##### {0} Discuss Answer #####\n{1}\n\n'.format(model1_name, model1_output_json))

    model2_thread.join()
    history_model2 = result_queue.get()
    model2_output_json = result_queue.get()
    model2_input_all = result_queue.get()
    logger.info("\n##### {0} Discuss Input_all #####\n{1}".format(model2_name, model2_input_all))
    logger.info('\n##### {0} Discuss Answer #####\n{1}\n\n'.format(model2_name, model2_output_json))
    add_para(log_doc, "\n##### {0} Discuss Input_all #####\n{1}".format(model2_name, model2_input_all))
    add_para_highlight(log_doc, '\n##### {0} Discuss Answer #####\n{1}\n\n'.format(model2_name, model2_output_json))

    model1_json = json.loads(model1_output_json.strip())
    model2_json = json.loads(model2_output_json.strip())
    model1_choose = model1_json["Superior solution"].strip()
    model2_choose = model2_json["Superior solution"].strip()
    logger.info("##### Choose Result ({0}, {1}) #####\n{2}: {3}\n{4}: {5}\n".format(
        path1[0], path2[0], model1_name, model1_json, model2_name, model2_json))
    add_para_highlight(log_doc, "##### Choose Result ({0}, {1}) #####\n{2}: {3}\n{4}: {5}\n".format(
        path1[0], path2[0], model1_name, model1_json, model2_name, model2_json))
    if model1_choose == model2_choose and model1_choose.strip() in [path1[0], path2[0]]:
        logger.info("##### Superior solution choice is consistent：{0}".format(model1_choose))
        add_para(log_doc, "##### Superior solution choice is consistent：{0}".format(model1_choose))
    else:
        logger.info("Choices are inconsistent, start a discussion ...")
        add_para(log_doc, "Choices are inconsistent, start a discussion ...")
        turn = 0
        while turn < 3 and model1_choose != model2_choose:
            turn += 1
            give_agent_opinion_to_model2 = open(auto_dir + "cross/give_agent_opinion.txt", 'r').read().format(
                opinion=model1_json["Reason"])
            history_model2, answer, _ = post_server(history_model2[:-1], system_m, give_agent_opinion_to_model2, False)
            history_model2, model2_output_json, model2_input = post_server(history_model2, system_m, solution_2_json, True)
            logger.info('\n##### {0} Discuss Input #####\n{1}\n\n'.format(model2_name, model2_input))
            logger.info('\n##### {0} Discuss Answer #####\n{1}\n\n'.format(model2_name, model2_output_json))
            add_para(log_doc, '\n##### {0} Discuss Input #####\n{1}\n\n'.format(model2_name, model2_input))
            add_para_highlight(log_doc, '\n##### {0} Discuss Answer #####\n{1}\n\n'.format(model2_name, model2_output_json))
            model2_json = json.loads(model2_output_json.strip())

            give_agent_opinion_to_model1 = open(auto_dir + "cross/give_agent_opinion.txt", 'r').read().format(
                opinion=model2_json["Reason"])
            history_model1, answer, input_all = dolphin_chat_utils.chat(
                history_model1[:-1], system_m, model1, tokenizer1, model_name, give_agent_opinion_to_model1, logger,
                model_tokenizer_device, False)
            history_model1, model1_output_json, model1_input = dolphin_chat_utils.chat(
                history_model1, system_m, model1, tokenizer1, model_name, solution_2_json, logger,
                model_tokenizer_device, True)
            logger.info('\n##### {0} Discuss Input #####\n{1}\n\n'.format(model1_name, model1_input))
            logger.info('\n##### {0} Discuss Answer #####\n{1}\n\n'.format(model1_name, model1_output_json))
            add_para(log_doc, '\n##### {0} Discuss Input #####\n{1}\n\n'.format(model1_name, model1_input))
            add_para_highlight(log_doc, '\n##### {0} Discuss Answer #####\n{1}\n\n'.format(model1_name, model1_output_json))
            model1_json = json.loads(model1_output_json.strip())
            model1_choose = model1_json["Superior solution"].strip()
            model2_choose = model2_json["Superior solution"].strip()

    if model1_choose == model2_choose:
        judge_choose = model1_choose
        logger.info("##### After discussion, the final choice is: {0}\n".format(judge_choose))
        add_para_highlight(log_doc, "##### After discussion, the final choice is: {0}\n".format(judge_choose))
    else:
        logger.info("The choices are still inconsistent after discussion, the final judgment is made ...")
        add_para(log_doc, "The choices are still inconsistent after discussion, the final judgment is made ...")
        history_judge = []
        role_path = "prompts/model_role/judge"
        system_m = open(role_path, 'r', encoding='utf-8').read().replace("\n", " ")
        solution_judge_1 = open(auto_dir + "cross/solution_judge_1.txt", 'r').read().format(
            Question=question.strip(), solution1=path1[1], solution2=path2[1], agent1=str(model1_json),
            agent2=str(model2_json))
        solution_judge_2 = open(auto_dir + "cross/solution_judge_2.txt", 'r').read()
        history_judge, answer, input_all = dolphin_chat_utils.chat(
            history_judge, system_m, model1, tokenizer1, model_name, solution_judge_1, logger, model_tokenizer_device,
            False)
        history_judge, judge_output_json, judge_input = dolphin_chat_utils.chat(
            history_judge, system_m, model1, tokenizer1, model_name, solution_judge_2, logger, model_tokenizer_device,
            True)
        logger.info('\n##### {0} Judge Input #####\n{1}\n\n'.format(model1_name, judge_input))
        logger.info('\n##### {0} Judge Answer #####\n{1}\n\n'.format(model1_name, judge_output_json))
        judge_json = json.loads(judge_output_json.strip())
        judge_choose = judge_json["Superior solution"].strip()
        logger.info("##### After judge, the final choice is: {0}\n".format(judge_choose))
        add_para(log_doc, '\n##### {0} Judge Input #####\n{1}\n\n'.format(model1_name, judge_input))
        add_para_highlight(log_doc, '\n##### {0} Judge Answer #####\n{1}\n\n'.format(model1_name, judge_output_json))
        add_para_highlight(log_doc, "##### After judge, the final choice is: {0}\n".format(judge_choose))
    if judge_choose == path1[0]:
        return path1
    elif judge_choose == path2[0]:
        return path2
    else:
        return None


def check_0_unreasonable(target_v, q_json_v, log_doc):
    role_path = "prompts/model_role/financial_analyst_assistant_sample"
    system_m = open(role_path, 'r', encoding='utf-8').read().replace("\n", " ")
    history_check = []
    check_0_prompt = open(auto_dir + "check/check_0.txt", 'r').read().format(t_v=target_v, q_json_v=q_json_v)
    check_0_json_prompt = open(auto_dir + "check/check_0_json.txt", 'r').read()

    history_check, _, _ = dolphin_chat_utils.chat(history_check, system_m, model1, tokenizer1, model1_name,
                                                  check_0_prompt, logger, model_tokenizer_device, False)
    history_check, check_0_output, check_0_input = dolphin_chat_utils.chat(history_check, system_m, model1, tokenizer1,
                                                                           model1_name, check_0_json_prompt, logger,
                                                                           model_tokenizer_device, True)
    logger.info('\n##### {0} Check 0 Input #####\n{1}\n\n'.format(model1_name, check_0_input))
    logger.info('\n##### {0} Check 0 Answer #####\n{1}\n\n'.format(model1_name, check_0_output))
    add_title(log_doc, "数值解为0的合理性检查")
    add_para(log_doc, '\n##### {0} Check 0 Input #####\n{1}\n\n'.format(model1_name, check_0_input))
    add_para_highlight(log_doc, '\n##### {0} Check 0 Answer #####\n{1}\n\n'.format(model1_name, check_0_output))

    check_0_json = json.loads(check_0_output.strip())
    check_0_choose = check_0_json["It is reasonable for this target variable to be zero"].strip()
    if check_0_choose == "No":
        return True
    else:
        return False


def check_negative_unreasonable(target_v, q_json_v, log_doc):
    role_path = "prompts/model_role/financial_analyst_assistant_sample"
    system_m = open(role_path, 'r', encoding='utf-8').read().replace("\n", " ")
    history_check = []
    check_symbol_prompt = open(auto_dir + "check/check_symbol.txt", 'r').read().format(t_v=target_v, q_json_v=q_json_v)
    check_symbol_json_prompt = open(auto_dir + "check/check_symbol_json.txt", 'r').read()

    history_check, _, _ = dolphin_chat_utils.chat(history_check, system_m, model1, tokenizer1, model1_name,
                                                  check_symbol_prompt, logger, model_tokenizer_device, False)
    history_check, check_symbol_output, check_symbol_input = dolphin_chat_utils.chat(history_check, system_m, model1,
                                                                                     tokenizer1, model1_name,
                                                                                     check_symbol_json_prompt, logger,
                                                                                     model_tokenizer_device, True)
    logger.info('\n##### {0} Check Symbol Input #####\n{1}\n\n'.format(model1_name, check_symbol_input))
    logger.info('\n##### {0} Check Symbol Answer #####\n{1}\n\n'.format(model1_name, check_symbol_output))
    add_title(log_doc, "数值解有正有负，数值解为负数的合理性检查")
    add_para(log_doc, '\n##### {0} Check Symbol Input #####\n{1}\n\n'.format(model1_name, check_symbol_input))
    add_para_highlight(log_doc, '\n##### {0} Check Symbol Answer #####\n{1}\n\n'.format(model1_name, check_symbol_output))

    check_symbol_json = json.loads(check_symbol_output.strip())
    check_symbol_choose = check_symbol_json["It is reasonable for this target variable to be negative"].strip()
    if check_symbol_choose == "No":
        return True
    else:
        return False


def contains_positive_and_negative(lst):
    has_positive = any(x > 0 for x in lst)
    has_negative = any(x < 0 for x in lst)
    return has_positive and has_negative


def post_server(history, mode_system_m, user_input, is_json):
    input_json = {"history": history, "system_m": mode_system_m, "user_input": user_input, "is_json": is_json}
    result = requests.post(dolphin_2_6_server + "post", json=json.dumps(input_json)).json()
    history = result["history"]
    answer = result["answer"]
    input_all = result["input_all"]
    return history, answer, input_all


if __name__ == '__main__':
    main()
