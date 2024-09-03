from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_title(doc, title_text):
    title = doc.add_paragraph()
    run = title.add_run(title_text)
    run.font.size = Pt(22)
    run.font.name = 'Times New Roman'
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_para(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_para_highlight(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)
    highlight = OxmlElement('w:highlight')
    highlight.set(qn('w:val'), 'lightGray')  # 设置高亮颜色
    run._element.rPr.append(highlight)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_dictionary(doc, dictionary, indent=0):
    indent_str = "    " * indent  # 定义缩进字符串
    for key, value in dictionary.items():
        if isinstance(value, dict):
            # 如果值是字典，则递归处理
            doc.add_paragraph(f"{indent_str}{key}:")
            add_dictionary(doc, value, indent + 1)
        elif isinstance(value, list):
            # 如果值是列表，遍历列表中的每个元素
            doc.add_paragraph(f"{indent_str}{key}:")
            for item in value:
                if isinstance(item, dict):
                    # 如果列表中的元素是字典，则递归处理
                    add_dictionary(doc, item, indent + 1)
                else:
                    # 如果列表中的元素是其他类型，直接添加
                    doc.add_paragraph(f"{indent_str}    - {item}")
        else:
            # 如果值不是字典也不是列表，直接添加
            doc.add_paragraph(f"{indent_str}{key}: {value}")